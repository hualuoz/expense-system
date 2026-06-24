#!/usr/bin/env python3
"""
文件解析模块 - 从截图、Word、PDF、Excel等文件中提取费用信息
"""

import io
import re
import json
from datetime import datetime

# ============ 费用类别映射 ============

CATEGORY_KEYWORDS = {
    "TRAVEL-TRANSPORT": ["交通", "高铁", "火车", "飞机", "机票", "动车", "航班", "车票", "地铁", "机票", "机票费"],
    "TRAVEL-LODGING": ["住宿", "酒店", "宾馆", "旅馆", "民宿", "房费"],
    "TRAVEL-MEAL": ["餐饮补贴", "餐补", "出差餐费", "伙食补贴"],
    "TRAVEL-OTHER": ["差旅其他", "行李", "签证", "差旅费"],
    "OFFICE-SUPPLIES": ["办公用品", "文具", "耗材", "打印纸", "墨盒", "笔", "文件夹"],
    "OFFICE-EQUIPMENT": ["办公设备", "电脑", "显示器", "键盘", "鼠标", "打印机"],
    "OFFICE-SOFTWARE": ["软件", "订阅", "SaaS", "会员", "云服务"],
    "OFFICE-PRINT": ["印刷", "名片", "宣传册", "海报"],
    "OFFICE-POST": ["快递", "邮寄", "物流", "邮费"],
    "ENTERTAIN-MEAL": ["招待", "业务餐", "宴请", "请客", "餐费"],
    "ENTERTAIN-GIFT": ["礼品", "伴手礼", "礼物"],
    "TRAINING-COURSE": ["培训", "课程", "讲座", "研讨"],
    "TRAINING-CERT": ["认证", "考试", "考证"],
    "COMM-MOBILE": ["话费", "通讯费", "手机费"],
    "COMM-NETWORK": ["网络", "宽带", "网费"],
    "TRANSPORT-TAXI": ["出租", "打车", "滴滴", "网约车"],
    "TRANSPORT-PARKING": ["停车", "停车费"],
    "TRANSPORT-FUEL": ["加油", "油费", "燃油", "汽油"],
    "TRANSPORT-PUBLIC": ["公交", "地铁卡", "充值"],
    "OTHER-MISC": ["杂费", "其他"],
}

CITY_KEYWORDS = {
    "北京": ["北京", "京城", "BJ"],
    "上海": ["上海", "沪", "SH"],
    "广州": ["广州", "穗", "GZ"],
    "深圳": ["深圳", "深", "SZ"],
    "成都": ["成都", "蓉", "CD"],
    "杭州": ["杭州", "杭", "HZ"],
    "武汉": ["武汉", "汉", "WH"],
    "南京": ["南京", "宁", "NJ"],
}


def classify_category(text):
    """根据文本内容自动分类费用类别"""
    text = text.upper()
    best_match = ("OTHER-MISC", 0)
    for category, keywords in CATEGORY_KEYWORDS.items():
        for kw in keywords:
            if kw in text:
                score = len(kw)  # 更长的关键词权重更高
                if score > best_match[1]:
                    best_match = (category, score)
    return best_match[0]


def extract_city(text):
    """从文本中提取城市名"""
    for city, keywords in CITY_KEYWORDS.items():
        for kw in keywords:
            if kw in text:
                return city
    return ""


def extract_amounts(text):
    """从文本中提取金额"""
    amounts = []
    # 匹配 ¥553.00 / ￥553 / 553.00元 等格式
    patterns = [
        r'[¥￥]\s*(\d+[,，]?\d*\.?\d*)',
        r'(\d+[,，]?\d*\.?\d*)\s*元',
        r'金额[：:]\s*(\d+[,，]?\d*\.?\d*)',
        r'合计[：:]\s*(\d+[,，]?\d*\.?\d*)',
        r'总计[：:]\s*(\d+[,，]?\d*\.?\d*)',
        r'费用[：:]\s*(\d+[,，]?\d*\.?\d*)',
        r'(\d+[,，]?\d*\.?\d{2})\s*(?:元|$)',  # 如 553.00
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            val = match.group(1).replace(',', '').replace('，', '')
            try:
                amounts.append(float(val))
            except ValueError:
                pass
    return amounts


def extract_dates(text):
    """从文本中提取日期"""
    dates = []
    patterns = [
        r'(\d{4})[年/\-.](\d{1,2})[月/\-.](\d{1,2})[日号]?',
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            y, m, d = match.group(1), match.group(2), match.group(3)
            try:
                date_str = f"{int(y):04d}-{int(m):02d}-{int(d):02d}"
                datetime.strptime(date_str, "%Y-%m-%d")
                dates.append(date_str)
            except ValueError:
                pass
    return dates


def extract_receipt_nos(text):
    """从文本中提取发票/票据号"""
    receipt_nos = []
    patterns = [
        r'(?:发票号|票据号|发票编号|No|编号)[：:\s]*(\w+[-]?\w+)',
        r'FP\d+',
        r'\d{8,20}',  # 长数字可能是发票号
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            val = match.group(0) if match.lastindex is None else match.group(1)
            receipt_nos.append(val)
    return receipt_nos


# ============ 文件解析器 ============

def parse_pdf(file_content):
    """解析 PDF 文件"""
    import pdfplumber
    results = []
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        full_text = ""
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text += text + "\n"
            # 尝试提取表格
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " ".join([cell or "" for cell in row])
                    full_text += row_text + "\n"

    expenses = _parse_text_to_expenses(full_text)
    return {
        "raw_text": full_text[:3000],
        "expenses": expenses,
        "file_type": "PDF",
    }


def parse_excel(file_content):
    """解析 Excel 文件"""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_content), data_only=True)
    results = []
    all_text = ""

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(cell) if cell is not None else "" for cell in row]
            row_text = " ".join(cells)
            all_text += row_text + "\n"

            if i == 0:
                headers = cells
                continue

            # 尝试将行数据映射为费用项
            expense = _map_row_to_expense(cells, headers)
            if expense and expense.get("amount"):
                results.append(expense)

    # 如果表格映射没成功，用文本解析兜底
    if not results:
        results = _parse_text_to_expenses(all_text)

    return {
        "raw_text": all_text[:3000],
        "expenses": results,
        "file_type": "Excel",
    }


def parse_word(file_content):
    """解析 Word 文件"""
    from docx import Document
    doc = Document(io.BytesIO(file_content))
    full_text = ""

    # 提取段落
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            full_text += " ".join(cells) + "\n"

    expenses = _parse_text_to_expenses(full_text)
    return {
        "raw_text": full_text[:3000],
        "expenses": expenses,
        "file_type": "Word",
    }


def parse_image(file_content):
    """解析图片文件（截图等）- 提取基础信息"""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(file_content))
        width, height = img.size

        # 尝试 OCR（如果安装了 pytesseract）
        try:
            import pytesseract
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        except (ImportError, Exception):
            text = f"[图片信息] 尺寸: {width}x{height}, 格式: {img.format or '未知'}"
            text += "\n\n⚠️ 未安装 Tesseract OCR，无法识别图片文字。"
            text += "\n安装方式: pip install pytesseract，并安装 Tesseract 引擎"
            text += "\n下载地址: https://github.com/UB-Mannheim/tesseract/wiki"

        expenses = _parse_text_to_expenses(text)
        return {
            "raw_text": text[:3000],
            "expenses": expenses,
            "file_type": "图片",
        }
    except Exception as e:
        return {
            "raw_text": f"图片解析失败: {str(e)}",
            "expenses": [],
            "file_type": "图片",
        }


def parse_csv(file_content):
    """解析 CSV 文件"""
    import csv
    reader = csv.DictReader(io.StringIO(file_content.decode('utf-8-sig', errors='ignore')))
    results = []
    all_text = ""
    headers = []

    for i, row in enumerate(reader):
        if i == 0:
            headers = list(row.keys())
        cells = list(row.values())
        row_text = " ".join(cells)
        all_text += row_text + "\n"

        expense = _map_row_to_expense(cells, headers)
        if expense and expense.get("amount"):
            results.append(expense)

    if not results:
        results = _parse_text_to_expenses(all_text)

    return {
        "raw_text": all_text[:3000],
        "expenses": results,
        "file_type": "CSV",
    }


# ============ 内部辅助函数 ============

# 中文表头映射到英文字段
HEADER_MAP = {
    "日期": "date", "时间": "date", "发生日期": "date", "费用日期": "date",
    "类别": "category", "费用类别": "category", "类型": "category", "分类": "category",
    "金额": "amount", "费用": "amount", "费用金额": "amount", "报销金额": "amount",
    "说明": "description", "描述": "description", "摘要": "description", "事由": "description", "用途": "description",
    "票据号": "receipt_no", "发票号": "receipt_no", "发票编号": "receipt_no", "单据号": "receipt_no",
    "城市": "city", "出差城市": "city",
    "级别": "employee_level", "员工级别": "employee_level",
    "项目编码": "project_code", "项目": "project_code",
}

AMOUNT_HEADERS = {"金额", "费用", "费用金额", "报销金额", "amount"}
DATE_HEADERS = {"日期", "时间", "发生日期", "费用日期", "date"}
DESC_HEADERS = {"说明", "描述", "摘要", "事由", "用途", "description"}


def _map_row_to_expense(cells, headers):
    """将表格行数据映射为费用项"""
    expense = {}
    mapped_headers = []

    # 先映射已知的表头
    for h in headers:
        h_stripped = h.strip() if h else ""
        eng = HEADER_MAP.get(h_stripped, h_stripped.lower().replace(" ", "_"))
        mapped_headers.append(eng)

    for i, cell in enumerate(cells):
        if i >= len(mapped_headers):
            break
        key = mapped_headers[i]
        val = str(cell).strip() if cell else ""
        if val:
            expense[key] = val

    # 转换金额为数字
    if "amount" in expense:
        try:
            expense["amount"] = float(str(expense["amount"]).replace(",", "").replace("，", "").replace("¥", "").replace("￥", ""))
        except (ValueError, TypeError):
            expense["amount"] = 0

    return expense


def _parse_text_to_expenses(text):
    """从纯文本中智能提取费用信息"""
    if not text or not text.strip():
        return []

    expenses = []
    amounts = extract_amounts(text)
    dates = extract_dates(text)
    category = classify_category(text)
    city = extract_city(text)
    receipt_nos = extract_receipt_nos(text)

    if not amounts:
        return []

    # 为每个金额创建一个费用项
    for i, amount in enumerate(amounts[:20]):  # 最多处理20笔
        expense = {
            "date": dates[i] if i < len(dates) else (dates[0] if dates else ""),
            "category": category,
            "amount": amount,
            "description": _extract_description(text, amount),
            "receipt_no": receipt_nos[i] if i < len(receipt_nos) else "",
            "city": city,
            "employee_level": "staff",
        }
        expenses.append(expense)

    return expenses


def _extract_description(text, amount):
    """提取金额附近的描述文字"""
    amount_str = str(amount)
    # 查找金额在文本中的位置
    for pattern in [f'¥{amount_str}', f'￥{amount_str}', f'{amount_str}元', amount_str]:
        idx = text.find(pattern)
        if idx >= 0:
            # 取金额前30个字符作为描述
            start = max(0, idx - 30)
            desc = text[start:idx].strip()
            # 清理换行符
            desc = re.sub(r'[\n\r]+', ' ', desc)
            if desc:
                return desc[:50]
    return ""


# ============ 统一入口 ============

def parse_file(filename, file_content):
    """
    统一文件解析入口
    :param filename: 文件名
    :param file_content: 文件二进制内容
    :return: 解析结果 dict
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    parsers = {
        "pdf": parse_pdf,
        "xlsx": parse_excel,
        "xls": parse_excel,
        "docx": parse_word,
        "doc": parse_word,
        "png": parse_image,
        "jpg": parse_image,
        "jpeg": parse_image,
        "gif": parse_image,
        "bmp": parse_image,
        "webp": parse_image,
        "csv": parse_csv,
    }

    parser = parsers.get(ext)
    if not parser:
        # 尝试作为纯文本解析
        try:
            text = file_content.decode("utf-8", errors="ignore")
            expenses = _parse_text_to_expenses(text)
            return {
                "raw_text": text[:3000],
                "expenses": expenses,
                "file_type": f"文本文件(.{ext})",
            }
        except Exception:
            return {
                "raw_text": f"不支持的文件格式: .{ext}",
                "expenses": [],
                "file_type": f"未知(.{ext})",
                "error": f"不支持的文件格式: .{ext}，支持格式: PDF, Excel, Word, 图片, CSV",
            }

    try:
        result = parser(file_content)
        result["filename"] = filename
        return result
    except Exception as e:
        return {
            "raw_text": f"解析失败: {str(e)}",
            "expenses": [],
            "file_type": ext.upper(),
            "filename": filename,
            "error": str(e),
        }
